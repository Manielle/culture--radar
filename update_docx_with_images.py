#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches
import os

# Ouvrir le document existant
doc = Document('/root/culture-radar/Culture_Radar_Cahier_Charges_Techniques.docx')

# Trouver l'endroit où insérer les images dans la section 7.a (après "Résultats des tris")
# On va parcourir les paragraphes pour trouver le bon endroit
paragraphs = doc.paragraphs

# Trouver l'index après "Conclusions" de la section 7.a
insert_index = None
for i, para in enumerate(paragraphs):
    if 'Conclusions' in para.text and i > 20:  # S'assurer qu'on est dans la section 7
        # Insérer après le paragraphe des conclusions de la section tri de cartes
        for j in range(i+1, len(paragraphs)):
            if 'task-oriented' in paragraphs[j].text:
                insert_index = j + 1
                break
        break

if insert_index:
    # Ajouter un nouveau paragraphe pour le titre
    new_para = paragraphs[insert_index-1].insert_paragraph_before('\nSchémas visuels des tris de cartes')
    new_para.style = 'Heading 3'
    
    # Ajouter l'image du tri de cartes
    p1 = paragraphs[insert_index-1].insert_paragraph_before('')
    run1 = p1.add_run()
    if os.path.exists('/root/culture-radar/tri_cartes_schema.png'):
        run1.add_picture('/root/culture-radar/tri_cartes_schema.png', width=Inches(6))
    
    # Ajouter l'image de la session
    p2 = paragraphs[insert_index-1].insert_paragraph_before('')
    run2 = p2.add_run()
    if os.path.exists('/root/culture-radar/session_tri_cartes.png'):
        run2.add_picture('/root/culture-radar/session_tri_cartes.png', width=Inches(5.5))

# Trouver l'endroit pour insérer l'arborescence (après la section arborescence)
for i, para in enumerate(paragraphs):
    if '└── Pages annexes' in para.text:
        # Ajouter après le dernier élément de l'arborescence
        new_para2 = para.insert_paragraph_before('\nSchéma visuel de l\'arborescence')
        new_para2.style = 'Heading 3'
        
        p3 = para.insert_paragraph_before('')
        run3 = p3.add_run()
        if os.path.exists('/root/culture-radar/arborescence_schema.png'):
            run3.add_picture('/root/culture-radar/arborescence_schema.png', width=Inches(6.5))
        break

# Sauvegarder le document mis à jour
doc.save('/root/culture-radar/Culture_Radar_Cahier_Charges_Techniques_Complet.docx')
print("✅ Document mis à jour avec les images : Culture_Radar_Cahier_Charges_Techniques_Complet.docx")