#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Créer un nouveau document
doc = Document()

# Fonction pour le style des titres
def set_heading_style(heading, level=1):
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        if level == 1:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(102, 126, 234)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(118, 75, 162)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True

# PAGE DE GARDE
title = doc.add_heading('CULTURE RADAR', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Manuel Utilisateur Complet')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(24)
subtitle.runs[0].font.bold = True

doc.add_paragraph('')
doc.add_paragraph('')

version = doc.add_paragraph('Version 1.0 - Août 2024')
version.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
doc.add_paragraph('')

tagline = doc.add_paragraph('Votre guide complet pour maîtriser la plateforme de découverte culturelle')
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.runs[0].font.italic = True
tagline.runs[0].font.size = Pt(14)

# INTRODUCTION
doc.add_page_break()
h1 = doc.add_heading('Introduction', 1)
set_heading_style(h1, 1)

doc.add_paragraph(
    'Bienvenue dans l\'univers de Culture Radar, votre nouvelle plateforme de référence pour la découverte '
    'd\'événements culturels personnalisés. Ce manuel a été conçu pour vous accompagner pas à pas dans '
    'l\'utilisation de notre plateforme, que vous soyez un nouvel utilisateur découvrant nos services ou '
    'un membre actif souhaitant approfondir sa maîtrise des fonctionnalités avancées.'
)

doc.add_paragraph(
    'Culture Radar représente une révolution dans la manière dont nous découvrons et vivons la culture. '
    'Née de la frustration de voir tant d\'événements passionnants rester invisibles au grand public, notre '
    'plateforme utilise l\'intelligence artificielle la plus avancée pour créer un pont entre vous et la '
    'richesse culturelle qui vous entoure. Chaque jour, des milliers d\'événements se déroulent dans votre '
    'ville, mais comment trouver ceux qui correspondent vraiment à vos goûts, vos disponibilités et vos envies '
    'du moment ? C\'est précisément la mission que s\'est donnée Culture Radar.'
)

doc.add_paragraph(
    'Notre approche est unique car elle ne se contente pas de lister des événements. Nous analysons en temps '
    'réel une multitude de facteurs : vos préférences personnelles bien sûr, mais aussi votre localisation, '
    'la météo, les conditions de transport, votre historique de participation, et même les tendances culturelles '
    'émergentes dans votre quartier. Cette analyse sophistiquée nous permet de vous présenter non pas une liste '
    'exhaustive et intimidante, mais une sélection soigneusement personnalisée d\'expériences culturelles qui '
    'ont toutes les chances de vous séduire.'
)

doc.add_paragraph(
    'Ce manuel est structuré pour faciliter votre apprentissage progressif. Les premiers chapitres vous '
    'guideront dans la création de votre compte et la découverte de l\'interface. Puis, nous explorerons '
    'ensemble les fonctionnalités de recherche et de recommandation qui font la force de Culture Radar. '
    'Les chapitres suivants détailleront les outils de gestion personnelle comme les favoris et le calendrier, '
    'avant d\'aborder les aspects sociaux et communautaires de la plateforme. Enfin, nous terminerons par '
    'les fonctionnalités premium et les ressources de support disponibles.'
)

# CHAPITRE 1 - DÉCOUVERTE DE CULTURE RADAR
doc.add_page_break()
h1 = doc.add_heading('Chapitre 1 : Découverte de Culture Radar', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('L\'histoire et la vision', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Culture Radar est né d\'une observation simple mais frustrante : dans une ville comme Paris, des centaines '
    'd\'événements culturels se déroulent chaque jour, mais la plupart restent invisibles pour le grand public. '
    'Isabelle Lemoine, notre fondatrice, a vécu cette frustration personnellement. Après quinze ans dans l\'industrie '
    'culturelle, elle connaissait l\'incroyable richesse de l\'offre culturelle française, mais constatait aussi '
    'le fossé énorme entre cette offre et sa visibilité auprès du public.'
)

doc.add_paragraph(
    'Le déclic est venu lors d\'une soirée entre amis où chacun se plaignait qu\'il ne se passait jamais rien '
    'd\'intéressant dans leur quartier. Isabelle savait que c\'était faux : elle venait de découvrir douze '
    'événements gratuits dans un rayon de deux kilomètres pour ce même week-end. Le problème n\'était pas '
    'l\'absence d\'événements, mais leur découvrabilité. Les plateformes existantes étaient soit trop généralistes, '
    'soit trop complexes, soit limitées à certains types d\'événements. Aucune n\'offrait cette expérience '
    'personnalisée et intuitive que nous attendons aujourd\'hui de nos services numériques.'
)

doc.add_paragraph(
    'C\'est ainsi qu\'est née l\'idée de Culture Radar : créer une plateforme qui pense à votre place, qui '
    'comprend vos goûts, vos contraintes et vos envies pour vous proposer exactement les événements qui vous '
    'correspondent. Nous ne voulions pas créer un énième agenda culturel, mais un véritable assistant personnel '
    'de découverte culturelle. Un service qui transforme la phrase "Je ne sais pas quoi faire ce soir" en '
    '"J\'ai l\'embarras du choix parmi ces suggestions parfaites".'
)

h2 = doc.add_heading('Les principes fondateurs', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Culture Radar repose sur quatre principes fondamentaux qui guident chaque décision de développement et '
    'chaque interaction avec nos utilisateurs. Le premier est l\'accessibilité universelle. Nous croyons '
    'fermement que la culture doit être accessible à tous, indépendamment de l\'âge, du niveau de revenus, '
    'de la localisation ou des capacités physiques. C\'est pourquoi notre plateforme est gratuite dans sa '
    'version de base, disponible sur tous les supports, et conçue pour être utilisable par les personnes '
    'en situation de handicap.'
)

doc.add_paragraph(
    'Le deuxième principe est la personnalisation intelligente. Contrairement aux plateformes qui vous '
    'noient sous des listes interminables, nous utilisons l\'intelligence artificielle pour comprendre '
    'précisément ce que vous recherchez. Notre algorithme apprend de vos interactions, s\'adapte à vos '
    'préférences changeantes et anticipe même vos envies en fonction du contexte. Si vous aimez le jazz '
    'mais que vous n\'avez jamais le temps en semaine, nous ne vous proposerons des concerts de jazz que '
    'pour vos soirées libres. Si vous préférez les activités en intérieur les jours de pluie, nous en '
    'tiendrons compte dans nos suggestions météo-sensibles.'
)

doc.add_paragraph(
    'Le troisième principe est la valorisation du local. Nous donnons autant d\'importance à la petite '
    'galerie de quartier qu\'au grand musée national. Notre mission est de révéler la richesse culturelle '
    'cachée de chaque territoire, de faire découvrir ces pépites locales qui font la vraie vie culturelle '
    'd\'une ville. Nous travaillons directement avec les organisateurs locaux, les associations, les '
    'collectivités pour garantir une couverture exhaustive de l\'offre culturelle.'
)

doc.add_paragraph(
    'Enfin, le quatrième principe est l\'éthique numérique. Nous respectons scrupuleusement votre vie '
    'privée, ne vendons jamais vos données, et sommes transparents sur notre utilisation de l\'IA. '
    'Vous gardez toujours le contrôle sur vos informations et pouvez à tout moment comprendre pourquoi '
    'nous vous avons fait telle ou telle recommandation.'
)

# CHAPITRE 2 - PREMIERS PAS
doc.add_page_break()
h1 = doc.add_heading('Chapitre 2 : Vos premiers pas sur la plateforme', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Accéder à Culture Radar', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Votre voyage culturel commence dès l\'instant où vous accédez à Culture Radar. Nous avons conçu '
    'plusieurs points d\'entrée pour nous adapter à vos habitudes numériques. La méthode la plus simple '
    'reste l\'accès via votre navigateur web préféré. En tapant www.culture-radar.fr dans la barre '
    'd\'adresse, vous accédez immédiatement à notre interface web responsive qui s\'adapte automatiquement '
    'à la taille de votre écran, que vous soyez sur ordinateur, tablette ou smartphone.'
)

doc.add_paragraph(
    'Pour une expérience optimale sur mobile, nous recommandons vivement notre application dédiée, '
    'disponible gratuitement sur l\'App Store pour les utilisateurs iOS et sur Google Play pour Android. '
    'L\'application mobile offre des avantages significatifs : notifications push pour ne jamais manquer '
    'un événement important, géolocalisation précise pour des suggestions ultra-pertinentes, mode hors '
    'ligne pour consulter vos favoris sans connexion, et intégration native avec votre calendrier et '
    'vos applications de partage préférées.'
)

doc.add_paragraph(
    'L\'installation de l\'application est un processus simple qui ne prend que quelques minutes. Sur iOS, '
    'ouvrez l\'App Store et recherchez "Culture Radar". Notre application apparaît avec notre logo distinctif '
    '- un radar stylisé aux couleurs violettes et bleues. Appuyez sur "Obtenir" et authentifiez-vous avec '
    'Face ID, Touch ID ou votre mot de passe Apple. L\'application, qui pèse environ 45 MB, se télécharge '
    'en quelques secondes avec une bonne connexion. Sur Android, le processus est similaire via le Google '
    'Play Store, avec une authentification via votre compte Google.'
)

h2 = doc.add_heading('Votre première visite', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Lors de votre première visite sur Culture Radar, nous avons préparé une expérience d\'accueil '
    'spécialement conçue pour vous familiariser avec la plateforme tout en commençant déjà à personnaliser '
    'votre expérience. Dès l\'arrivée sur la page d\'accueil, vous êtes accueilli par une interface épurée '
    'qui met immédiatement en avant ce qui compte : les événements culturels près de vous.'
)

doc.add_paragraph(
    'Même sans créer de compte, vous pouvez déjà explorer une sélection d\'événements populaires dans votre '
    'région. La plateforme détecte automatiquement votre ville (avec votre permission) et affiche les '
    'événements du jour et des prochains jours. Cette approche "try before you buy" vous permet de '
    'comprendre immédiatement la valeur de Culture Radar sans engagement. Vous pouvez parcourir les '
    'événements, lire leurs descriptions détaillées, voir les photos, consulter les informations pratiques '
    'comme les horaires et les tarifs.'
)

doc.add_paragraph(
    'Un bandeau discret en haut de page vous invite à créer un compte gratuit pour débloquer toutes les '
    'fonctionnalités. Ce n\'est pas obligatoire pour commencer, mais fortement recommandé pour profiter '
    'pleinement de l\'expérience Culture Radar. Le processus d\'inscription, que nous détaillerons dans '
    'le chapitre suivant, a été optimisé pour être aussi rapide et fluide que possible. En trois minutes, '
    'vous pouvez créer votre compte, définir vos préférences de base et commencer à recevoir des '
    'recommandations personnalisées.'
)

# CHAPITRE 3 - CRÉATION ET CONFIGURATION DU COMPTE
doc.add_page_break()
h1 = doc.add_heading('Chapitre 3 : Création et configuration de votre compte', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Le processus d\'inscription détaillé', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'La création de votre compte Culture Radar est une étape cruciale qui détermine la qualité de votre '
    'expérience sur la plateforme. Nous avons conçu ce processus pour être à la fois simple et complet, '
    'recueillant juste assez d\'informations pour personnaliser votre expérience sans être intrusif. '
    'Commençons par cliquer sur le bouton "S\'inscrire" situé en haut à droite de l\'écran. Ce bouton, '
    'volontairement mis en évidence avec notre couleur signature violet, ouvre un formulaire d\'inscription '
    'moderne et accueillant.'
)

doc.add_paragraph(
    'Vous avez le choix entre plusieurs méthodes d\'inscription, chacune ayant ses avantages. La méthode '
    'classique par email vous permet de créer un compte indépendant avec une adresse email et un mot de '
    'passe de votre choix. C\'est l\'option privilégiée si vous préférez garder vos comptes séparés ou '
    'si vous utilisez une adresse email dédiée aux activités culturelles. Le formulaire vous demande '
    'votre adresse email, un mot de passe sécurisé (nous y reviendrons), votre prénom pour personnaliser '
    'nos communications, et votre code postal pour localiser les événements près de vous.'
)

doc.add_paragraph(
    'Les options de connexion sociale via Google, Facebook ou Apple ID offrent une inscription en un clic. '
    'Ces méthodes sont particulièrement pratiques car elles évitent la création d\'un nouveau mot de passe '
    'et permettent une connexion ultérieure très rapide. Lorsque vous choisissez l\'une de ces options, '
    'nous récupérons uniquement les informations de base nécessaires : votre nom, votre email, et votre '
    'photo de profil si vous le souhaitez. Nous ne publions jamais rien sur vos réseaux sociaux sans '
    'votre autorisation explicite, et nous n\'accédons pas à votre liste d\'amis ou à vos publications.'
)

doc.add_paragraph(
    'Concernant la sécurité du mot de passe, si vous optez pour l\'inscription classique, nous appliquons '
    'des standards de sécurité stricts. Votre mot de passe doit contenir au minimum huit caractères, '
    'incluant au moins une majuscule, une minuscule, un chiffre et un caractère spécial. Cette complexité '
    'est nécessaire pour protéger votre compte contre les tentatives d\'intrusion. Notre système vérifie '
    'également que votre mot de passe n\'apparaît pas dans les listes de mots de passe compromis connus. '
    'Un indicateur de force en temps réel vous guide pour créer un mot de passe robuste.'
)

h2 = doc.add_heading('Configuration initiale du profil', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Une fois votre compte créé, vous êtes guidé à travers un processus de configuration initiale qui '
    'détermine la pertinence de nos futures recommandations. Cette étape, bien que facultative, est '
    'fortement recommandée car elle améliore considérablement la qualité de votre expérience. Le processus '
    'se déroule en trois écrans successifs, chacun se concentrant sur un aspect différent de vos préférences.'
)

doc.add_paragraph(
    'Le premier écran vous invite à sélectionner vos catégories culturelles préférées. Présenté sous forme '
    'de cartes visuelles attractives, vous pouvez choisir parmi : Concerts & Musique, Théâtre & Spectacles, '
    'Expositions & Musées, Cinéma, Conférences & Débats, Ateliers & Cours, Festivals, et Sport & Loisirs. '
    'Pour chaque catégorie, vous pouvez affiner en indiquant votre niveau d\'intérêt sur une échelle de '
    'un à cinq étoiles. Cette granularité nous permet de comprendre non seulement ce que vous aimez, mais '
    'aussi l\'intensité de votre intérêt. Par exemple, vous pourriez mettre cinq étoiles aux concerts de '
    'jazz mais seulement trois étoiles au théâtre classique.'
)

doc.add_paragraph(
    'Le deuxième écran concerne vos préférences pratiques. Ici, vous définissez votre budget habituel par '
    'sortie (gratuit, moins de 20€, 20-50€, plus de 50€, ou sans limite), la distance maximale que vous '
    'êtes prêt à parcourir (1km pour le quartier, 5km pour la ville, 10km pour l\'agglomération, ou sans '
    'limite), vos jours de sortie préférés (semaine, week-end, ou indifférent), et vos créneaux horaires '
    'privilégiés (matinée, après-midi, soirée, ou nocturne). Ces informations sont cruciales pour filtrer '
    'les milliers d\'événements disponibles et ne vous présenter que ceux qui correspondent à vos contraintes '
    'pratiques.'
)

doc.add_paragraph(
    'Le troisième et dernier écran configure vos préférences de communication. Vous choisissez comment et '
    'quand recevoir nos recommandations : notifications push sur mobile, emails récapitulatifs, ou les deux. '
    'Vous pouvez opter pour des alertes en temps réel pour les événements de dernière minute, un digest '
    'quotidien le matin avec les événements du jour, ou un récapitulatif hebdomadaire le jeudi pour '
    'planifier votre week-end. Vous définissez également vos heures de tranquillité où vous ne souhaitez '
    'recevoir aucune notification, typiquement pendant vos heures de travail ou la nuit.'
)

# CHAPITRE 4 - NAVIGATION ET INTERFACE
doc.add_page_break()
h1 = doc.add_heading('Chapitre 4 : Maîtriser l\'interface et la navigation', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Architecture de l\'interface', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'L\'interface de Culture Radar a été méticuleusement conçue selon les principes du design moderne et '
    'de l\'ergonomie cognitive. Chaque élément a sa place, chaque interaction est intuitive, et l\'ensemble '
    'crée une expérience fluide qui disparaît au profit du contenu : les événements culturels. Commençons '
    'par explorer la structure générale de l\'interface, qui reste cohérente que vous soyez sur ordinateur, '
    'tablette ou smartphone, avec des adaptations intelligentes selon la taille de l\'écran.'
)

doc.add_paragraph(
    'La barre de navigation supérieure est l\'élément permanent qui vous accompagne sur toutes les pages. '
    'Elle contient les éléments essentiels pour naviguer rapidement. À gauche, le logo Culture Radar sert '
    'de point d\'ancrage visuel et de raccourci vers la page d\'accueil. Au centre, la barre de recherche '
    'intelligente occupe une position privilégiée, toujours accessible pour lancer une recherche rapide. '
    'À droite, vous trouvez les icônes d\'accès rapide : le cœur pour vos favoris, la cloche pour les '
    'notifications (avec un badge rouge indiquant le nombre de nouvelles notifications), et votre avatar '
    'pour accéder au menu utilisateur.'
)

doc.add_paragraph(
    'Le menu latéral gauche, rétractable sur mobile pour gagner de l\'espace, organise les principales '
    'sections de la plateforme. Chaque item du menu est accompagné d\'une icône distinctive et d\'un label '
    'clair. "Découvrir" vous emmène vers l\'exploration libre des événements, "Pour vous" affiche vos '
    'recommandations personnalisées, "Calendrier" présente vos événements planifiés dans une vue temporelle, '
    '"Favoris" rassemble vos événements sauvegardés, "Carte" offre une vue géographique des événements, '
    'et "Paramètres" vous permet de configurer votre expérience.'
)

doc.add_paragraph(
    'La zone centrale de contenu s\'adapte dynamiquement selon la section visitée. Sur la page d\'accueil, '
    'elle présente un mélange savamment orchestré de recommandations personnalisées, d\'événements tendances, '
    'et de découvertes du jour. Les événements sont présentés sous forme de cartes visuelles qui révèlent '
    'progressivement plus d\'informations selon votre niveau d\'interaction : au survol, la carte s\'anime '
    'subtilement et révèle des informations supplémentaires ; au clic, elle s\'ouvre dans une vue détaillée.'
)

h2 = doc.add_heading('Les cartes d\'événements', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Les cartes d\'événements sont le cœur de l\'expérience Culture Radar. Chaque carte est une invitation '
    'visuelle à découvrir un événement, conçue pour transmettre un maximum d\'informations de manière claire '
    'et attractive. La photo de couverture occupe les deux tiers supérieurs de la carte, soigneusement '
    'sélectionnée pour représenter l\'essence de l\'événement. Un dégradé subtil en bas de l\'image assure '
    'la lisibilité du texte superposé.'
)

doc.add_paragraph(
    'Le titre de l\'événement est affiché en caractères gras, optimisé pour être lu rapidement même en '
    'scrollant. Juste en dessous, une ligne de métadonnées essentielles : la date et l\'heure en premier, '
    'suivies du lieu avec la distance depuis votre position, et enfin le prix ou la mention "Gratuit" '
    'mise en évidence en vert. Un système de badges visuels indique rapidement les caractéristiques '
    'importantes : "Dernières places", "Nouveau", "Coup de cœur", "Accessible PMR", etc.'
)

doc.add_paragraph(
    'La partie inférieure de la carte contient les actions rapides accessibles sans ouvrir la fiche '
    'complète. Le bouton cœur permet d\'ajouter instantanément aux favoris, avec une animation satisfaisante '
    'qui confirme l\'action. Le bouton calendrier ajoute l\'événement à votre planning personnel. Le bouton '
    'partage ouvre un menu contextuel avec vos options de partage préférées. Ces actions sont conçues pour '
    'être effectuées d\'un seul geste, particulièrement important sur mobile où chaque tap compte.'
)

# CHAPITRE 5 - SYSTÈME DE RECHERCHE
doc.add_page_break()
h1 = doc.add_heading('Chapitre 5 : Le système de recherche intelligent', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Comprendre la recherche intelligente', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Le système de recherche de Culture Radar représente une avancée majeure dans la découverte d\'événements '
    'culturels. Contrairement aux moteurs de recherche traditionnels qui se contentent de chercher des mots-clés '
    'dans une base de données, notre système comprend véritablement votre intention. Quand vous tapez "quelque '
    'chose de sympa ce soir", le système analyse cette requête naturelle et comprend que vous cherchez un '
    'événement pour aujourd\'hui, en soirée, probablement divertissant et accessible.'
)

doc.add_paragraph(
    'Cette intelligence repose sur plusieurs couches technologiques sophistiquées. La première est le '
    'traitement du langage naturel qui décompose votre requête en intentions et en entités. "Concert de '
    'jazz demain soir pas cher" devient : type=concert, genre=jazz, date=demain, période=soir, budget=économique. '
    'La deuxième couche est contextuelle : elle prend en compte votre localisation actuelle, vos préférences '
    'enregistrées, votre historique de recherche, et même des facteurs externes comme la météo.'
)

doc.add_paragraph(
    'La troisième couche est prédictive. Le système apprend de vos comportements passés pour anticiper '
    'vos besoins. Si vous cherchez régulièrement des expositions le dimanche après-midi, le système '
    'privilégiera ce type de résultats quand vous ferez une recherche vague un dimanche matin. Si vous '
    'avez tendance à préférer les événements gratuits en fin de mois, cette préférence temporelle sera '
    'prise en compte automatiquement.'
)

doc.add_paragraph(
    'Au-delà de la recherche textuelle, Culture Radar innove avec la recherche vocale et visuelle. '
    'La recherche vocale, particulièrement pratique en mobilité, comprend les requêtes parlées naturellement : '
    '"Montre-moi ce qu\'il y a comme concerts ce week-end" fonctionne aussi bien qu\'une requête tapée. '
    'La recherche visuelle permet de scanner une affiche dans la rue pour obtenir instantanément les '
    'informations sur l\'événement et l\'ajouter à vos favoris.'
)

h2 = doc.add_heading('Utilisation avancée des filtres', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Les filtres de Culture Radar vont bien au-delà des options basiques de catégorie et de date. Nous '
    'avons développé un système de filtrage multidimensionnel qui permet d\'affiner votre recherche selon '
    'des critères très précis tout en restant simple d\'utilisation. Le panneau de filtres, accessible '
    'via l\'icône d\'entonnoir à côté de la barre de recherche, s\'ouvre en overlay sur desktop ou en '
    'bottom sheet sur mobile.'
)

doc.add_paragraph(
    'Les filtres temporels offrent une granularité remarquable. Au-delà des options évidentes comme '
    '"Aujourd\'hui", "Ce week-end" ou "Cette semaine", vous pouvez définir des créneaux précis. Le filtre '
    '"Après le travail" sélectionne automatiquement les événements entre 18h et 20h en semaine. "Matinée '
    'en famille" privilégie les événements du matin adaptés aux enfants. "Nocturne" met en avant les '
    'événements après 21h. Ces filtres intelligents s\'adaptent même à vos habitudes : si vous avez '
    'indiqué dans votre profil que vous travaillez en horaires décalés, "Après le travail" s\'ajustera '
    'en conséquence.'
)

doc.add_paragraph(
    'Les filtres géographiques vont au-delà du simple rayon kilométrique. Vous pouvez filtrer par quartier, '
    'par ligne de métro, ou même par temps de trajet. Le filtre "Sur mon trajet" identifie les événements '
    'situés entre votre domicile et votre lieu de travail, parfaits pour une sortie improvisée en rentrant. '
    'Le filtre "Accessible en transport" exclut les lieux difficiles d\'accès sans voiture. "Quartiers que '
    'j\'aime" se base sur vos lieux favoris précédents pour suggérer des zones similaires.'
)

doc.add_paragraph(
    'Les filtres thématiques permettent une recherche par ambiance ou par contexte plutôt que par catégorie '
    'stricte. "Romantique" sélectionne des événements propices aux rendez-vous amoureux, tous types confondus. '
    '"Intellectuel" privilégie les conférences, débats et expositions pointues. "Festif" met en avant les '
    'événements à forte énergie sociale. "Contempla' + 'tif" suggère des expériences calmes et introspectives. '
    'Ces filtres transversaux permettent de découvrir des événements auxquels vous n\'auriez pas pensé mais '
    'qui correspondent parfaitement à votre humeur du moment.'
)

# CHAPITRE 6 - RECOMMANDATIONS IA
doc.add_page_break()
h1 = doc.add_heading('Chapitre 6 : L\'intelligence artificielle au service de vos découvertes', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Le moteur de recommandation expliqué', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Le cœur battant de Culture Radar est son moteur de recommandation alimenté par intelligence artificielle, '
    'un système complexe que nous avons rendu invisible pour vous offrir une expérience magique. Imaginez '
    'un ami qui connaîtrait parfaitement vos goûts, suivrait l\'actualité culturelle en permanence, '
    'connaîtrait chaque recoin de votre ville, et pourrait instantanément croiser toutes ces informations '
    'pour vous faire LA suggestion parfaite. C\'est exactement ce que fait notre IA, mais à une échelle '
    'et avec une précision qu\'aucun humain ne pourrait atteindre.'
)

doc.add_paragraph(
    'Notre algorithme analyse en permanence des millions de points de données pour créer votre profil '
    'culturel unique. Ce profil n\'est pas statique : il évolue constamment en fonction de vos interactions '
    'avec la plateforme. Chaque événement que vous consultez, chaque favori que vous ajoutez, chaque '
    'recherche que vous effectuez enrichit votre profil. Mais nous allons plus loin que le simple historique. '
    'Le temps que vous passez à regarder une fiche événement, le moment où vous la consultez, si vous '
    'revenez plusieurs fois sur le même événement, si vous partagez l\'information avec des amis - tous '
    'ces micro-comportements sont analysés pour affiner notre compréhension de vos préférences.'
)

doc.add_paragraph(
    'La dimension temporelle est cruciale dans nos recommandations. Nous avons découvert que les goûts '
    'culturels ne sont pas constants mais fluctuent selon de nombreux facteurs. Vos envies du vendredi '
    'soir ne sont pas les mêmes que celles du dimanche après-midi. En hiver, vous pourriez privilégier '
    'les activités en intérieur, tandis qu\'en été, les festivals en plein air prennent le dessus. Après '
    'une semaine stressante, vous pourriez rechercher des événements relaxants, alors qu\'en période de '
    'vacances, vous êtes plus ouvert à l\'aventure et à la découverte. Notre IA intègre tous ces patterns '
    'temporels pour faire des recommandations contextuellement pertinentes.'
)

doc.add_paragraph(
    'La composante sociale enrichit considérablement nos recommandations. Si vos amis sur Culture Radar '
    'ont aimé un événement, il a plus de chances de vous plaire aussi. Mais nous ne nous contentons pas '
    'de cette approche simpliste. Nous analysons les réseaux d\'affinités : les personnes qui ont des '
    'goûts similaires aux vôtres sur certains aspects peuvent vous faire découvrir de nouveaux horizons '
    'sur d\'autres. C\'est ainsi que notre système peut vous surprendre avec des recommandations inattendues '
    'mais pertinentes, vous sortant doucement de votre zone de confort culturel.'
)

h2 = doc.add_heading('Les différents types de recommandations', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Culture Radar propose plusieurs types de recommandations, chacune répondant à un besoin spécifique '
    'de découverte culturelle. Les "Découvertes du jour" sont notre sélection quotidienne personnalisée, '
    'fraîchement calculée chaque matin à 7h. Ces cinq événements sont choisis non seulement pour leur '
    'adéquation avec vos goûts, mais aussi pour leur diversité. Nous veillons à mixer les types '
    'd\'événements, les lieux, les horaires et les budgets pour vous offrir un éventail de possibilités. '
    'Cette sélection apparaît en premier sur votre page d\'accueil et fait l\'objet d\'une notification '
    'push si vous l\'avez activée.'
)

doc.add_paragraph(
    'Les "Tendances de votre quartier" mettent en lumière ce qui fait vibrer votre environnement immédiat. '
    'Ces recommandations sont calculées en analysant l\'activité des utilisateurs dans votre zone '
    'géographique, pondérée par la similarité de leurs profils avec le vôtre. C\'est une excellente '
    'façon de découvrir les événements qui créent le buzz localement, ces pépites de quartier qui ne '
    'font pas forcément la une des médias mais qui créent du lien social et de la vie culturelle locale. '
    'Un café-concert intimiste, une exposition dans une galerie associative, une lecture publique dans '
    'une librairie - ces événements de proximité sont souvent les plus mémorables.'
)

doc.add_paragraph(
    'Les recommandations "Dernière chance" ont un caractère d\'urgence qui les rend particulièrement '
    'précieuses. Ce sont des événements qui correspondent parfaitement à vos goûts mais qui se terminent '
    'bientôt ou dont les dernières places sont disponibles. Notre système surveille en temps réel la '
    'disponibilité et les dates de fin pour vous alerter avant qu\'il ne soit trop tard. Ces alertes '
    'peuvent être paramétrées selon votre réactivité : certains utilisateurs veulent être prévenus une '
    'semaine avant, d\'autres préfèrent l\'adrénaline de la dernière minute.'
)

doc.add_paragraph(
    'Les "Chemins de découverte" représentent notre fonctionnalité la plus innovante en matière de '
    'recommandation. Plutôt que de vous enfermer dans vos préférences habituelles, nous créons des '
    'parcours progressifs pour vous faire découvrir de nouveaux horizons culturels. Si vous aimez le '
    'rock indépendant, nous pourrions vous suggérer un concert de jazz fusion, puis progressivement '
    'vous amener vers le jazz traditionnel. Ces chemins sont conçus pour respecter votre rythme de '
    'découverte, avec toujours un pied dans le familier et un pied dans la nouveauté.'
)

# CHAPITRE 7 - GESTION DES FAVORIS
doc.add_page_break()
h1 = doc.add_heading('Chapitre 7 : Organiser vos découvertes avec les favoris', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('La philosophie des favoris', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Les favoris sur Culture Radar sont bien plus qu\'une simple liste de marque-pages. Nous les concevons '
    'comme votre bibliothèque culturelle personnelle, un espace où vous construisez votre identité culturelle '
    'en collectionnant les expériences qui vous attirent. Chaque événement ajouté aux favoris raconte '
    'quelque chose sur vous : vos aspirations, vos curiosités, vos projets. Notre système de favoris a été '
    'pensé pour transformer cette collection passive en un outil actif de planification et de partage culturel.'
)

doc.add_paragraph(
    'L\'acte d\'ajouter un événement aux favoris est volontairement simple et satisfaisant. Un simple tap '
    'sur l\'icône cœur déclenche une micro-animation où le cœur se remplit et pulse légèrement, accompagné '
    'd\'une vibration haptique sur mobile. Cette feedback immédiat crée une satisfaction psychologique qui '
    'encourage l\'engagement. Mais derrière cette simplicité apparente se cache un système sophistiqué. '
    'L\'événement n\'est pas simplement stocké : il est analysé, catégorisé, et connecté à votre graphe '
    'de préférences pour améliorer futures recommandations.'
)

doc.add_paragraph(
    'Vos favoris sont synchronisés en temps réel sur tous vos appareils. Ajoutez un événement depuis votre '
    'ordinateur au bureau, et il apparaît instantanément sur votre téléphone. Cette synchronisation va '
    'au-delà du simple mirroring : chaque appareil adapte l\'affichage à son contexte. Sur mobile, les '
    'favoris sont optimisés pour la consultation rapide en déplacement. Sur tablette, ils s\'affichent '
    'dans une vue magazine riche visuellement. Sur desktop, vous accédez à tous les outils d\'organisation '
    'et de gestion avancés.'
)

h2 = doc.add_heading('Création et gestion des listes', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Les listes thématiques transforment vos favoris d\'une accumulation désordonnée en collections '
    'organisées et significatives. Créer une liste est un acte créatif où vous définissez non seulement '
    'un titre et une description, mais aussi une intention. La liste "Sorties romantiques" n\'est pas '
    'qu\'un dossier, c\'est un projet de moments à partager. La liste "Formation continue" représente '
    'votre volonté de croissance personnelle. "Activités avec les enfants" matérialise votre rôle parental.'
)

doc.add_paragraph(
    'Chaque liste peut être personnalisée avec une couleur et une icône, créant un système visuel qui '
    'facilite la navigation rapide. Mais les options vont plus loin. Vous pouvez définir des collaborateurs '
    'pour une liste, transformant la planification culturelle en activité sociale. Imaginez préparer un '
    'anniversaire surprise en créant une liste collaborative où chaque ami peut suggérer des événements. '
    'Ou un couple qui maintient une liste commune de sorties à faire, chacun alimentant les idées au fil '
    'de ses découvertes.'
)

doc.add_paragraph(
    'Les listes peuvent être rendues publiques, devenant alors des playlists culturelles que d\'autres '
    'utilisateurs peuvent suivre. Les conservateurs de musées, les critiques culturels, les influenceurs '
    'locaux créent des listes curatoriales suivies par des milliers de personnes. Vous-même pouvez devenir '
    'un prescripteur culturel en partageant vos sélections. Le système de likes et de commentaires sur les '
    'listes crée une dimension sociale enrichissante, où la découverte culturelle devient conversation.'
)

doc.add_paragraph(
    'La fonctionnalité "Smart Lists" pousse l\'organisation encore plus loin avec des listes dynamiques '
    'qui se mettent à jour automatiquement selon des critères définis. Créez une smart list "Concerts '
    'gratuits ce mois-ci" et elle se remplira automatiquement des événements correspondants. "Expositions '
    'finissant bientôt" vous rappelle ce qu\'il ne faut pas manquer. "Événements où vont mes amis" '
    'agrège l\'activité sociale de votre réseau. Ces listes intelligentes transforment la gestion des '
    'favoris de tâche manuelle en curation automatisée.'
)

# CHAPITRE 8 - CALENDRIER ET PLANIFICATION
doc.add_page_break()
h1 = doc.add_heading('Chapitre 8 : Votre calendrier culturel personnel', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Vue d\'ensemble du système de calendrier', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Le calendrier Culture Radar réinvente la planification culturelle en intégrant seamlessly vos envies '
    'culturelles dans votre vie quotidienne. Plus qu\'un simple agenda, c\'est un assistant personnel qui '
    'comprend vos contraintes temporelles et optimise vos opportunités de sorties. La vue calendrier, '
    'accessible depuis le menu principal ou via le bouton calendrier sur chaque événement, présente une '
    'interface familière mais augmentée de fonctionnalités spécifiques à la découverte culturelle.'
)

doc.add_paragraph(
    'La vue mensuelle offre une perspective globale de votre vie culturelle. Chaque jour affiche des '
    'pastilles colorées représentant vos événements planifiés, avec un code couleur intuitif : violet '
    'pour les concerts, bleu pour les expositions, vert pour les ateliers, orange pour le théâtre. '
    'Les jours avec des suggestions particulièrement pertinentes sont subtilement mis en évidence, vous '
    'invitant à explorer les possibilités. Un indicateur de "densité culturelle" montre les jours où '
    'votre ville bouillonne d\'activités, parfait pour planifier une sortie spontanée.'
)

doc.add_paragraph(
    'La vue hebdomadaire est votre cockpit de planification tactique. Elle affiche non seulement vos '
    'événements confirmés, mais aussi les "possibilités" - ces événements dans vos favoris qui pourraient '
    's\'intégrer dans votre planning. L\'intelligence artificielle analyse vos créneaux libres, vos '
    'habitudes de déplacement, et même votre niveau d\'énergie probable (basé sur votre historique) pour '
    'suggérer les meilleurs moments pour chaque sortie. Le vendredi après une semaine chargée ? Peut-être '
    'pas le meilleur moment pour une conférence de trois heures sur la physique quantique.'
)

doc.add_paragraph(
    'La vue liste chronologique est parfaite pour avoir une vision linéaire de vos prochaines sorties. '
    'Chaque événement est affiché avec toutes les informations pratiques : horaires, lieu avec temps de '
    'trajet estimé depuis votre position habituelle à cette heure, météo prévue, et même un rappel de '
    'avec qui vous y allez si vous avez partagé l\'événement. Des actions rapides permettent de modifier, '
    'annuler, ou inviter des amis directement depuis cette vue.'
)

h2 = doc.add_heading('Intégration et synchronisation', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'L\'intégration avec vos calendriers existants était une priorité absolue dans notre conception. '
    'Culture Radar ne veut pas remplacer votre système de gestion du temps, mais l\'enrichir. La '
    'synchronisation bidirectionnelle avec Google Calendar, Apple Calendar, et Outlook permet à vos '
    'événements culturels d\'apparaître dans votre calendrier principal, et inversement, vos engagements '
    'existants sont pris en compte pour éviter les conflits de planning.'
)

doc.add_paragraph(
    'La configuration de la synchronisation offre une granularité fine. Vous pouvez choisir de synchroniser '
    'tous vos événements ou seulement ceux marqués comme "confirmés". Les événements peuvent être créés '
    'dans un calendrier dédié "Culture Radar" ou intégrés à votre calendrier principal. Les informations '
    'synchronisées incluent non seulement les basics (titre, date, lieu) mais aussi les liens vers la '
    'billetterie, les notes personnelles, et même un lien de retour vers Culture Radar pour accéder aux '
    'informations complètes et aux fonctionnalités sociales.'
)

doc.add_paragraph(
    'L\'import de votre calendrier existant permet à Culture Radar de devenir intelligent dès le premier '
    'jour. En analysant vos engagements actuels, notre IA comprend vos patterns de disponibilité. Elle '
    'remarque que vous avez yoga tous les mardis soirs, que vos weekends pairs sont généralement pris, '
    'que vous préférez garder les dimanches matins libres. Cette compréhension permet des recommandations '
    'qui s\'intègrent naturellement dans votre vie plutôt que de la perturber.'
)

doc.add_paragraph(
    'La gestion des conflits et des changements est particulièrement soignée. Si un événement est annulé '
    'ou reporté, vous êtes immédiatement notifié avec des propositions d\'alternatives. Si deux événements '
    'se chevauchent, l\'interface vous aide à arbitrer en présentant les pour et contre de chaque option : '
    'distance, coût, rareté de l\'événement, présence d\'amis. Le système apprend même de vos arbitrages '
    'passés pour mieux vous conseiller à l\'avenir.'
)

# CHAPITRE 9 - ASPECT SOCIAL
doc.add_page_break()
h1 = doc.add_heading('Chapitre 9 : La dimension sociale de Culture Radar', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Construire sa communauté culturelle', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Culture Radar transforme la découverte culturelle solitaire en expérience sociale enrichissante. '
    'La culture a toujours été un vecteur de lien social, et notre plateforme amplifie cette dimension '
    'en facilitant le partage, la découverte commune et les sorties de groupe. Votre réseau sur Culture '
    'Radar n\'est pas juste une liste d\'amis, c\'est un écosystème vivant d\'échanges culturels où '
    'chaque membre enrichit l\'expérience des autres.'
)

doc.add_paragraph(
    'Construire votre réseau commence naturellement. Lors de l\'inscription, vous pouvez connecter vos '
    'comptes sociaux pour retrouver vos amis déjà présents sur la plateforme. Mais nous allons plus loin '
    'avec notre système de "découverte d\'affinités" qui suggère des connexions basées sur les goûts '
    'culturels partagés. Vous adorez le théâtre expérimental et le jazz manouche ? Nous vous présenterons '
    'd\'autres passionnés de ces niches culturelles. Ces connexions par affinité créent souvent des '
    'amitiés plus profondes que les réseaux sociaux traditionnels car elles sont basées sur des passions '
    'partagées.'
)

doc.add_paragraph(
    'Le feed social de Culture Radar est soigneusement curé pour rester pertinent et inspirant. Plutôt '
    'qu\'un flux chronologique simple, nous priorisons les activités qui vous intéressent vraiment : un '
    'ami qui va à un événement que vous avez en favori, quelqu\'un de votre réseau qui découvre un lieu '
    'que vous ne connaissez pas, une review détaillée d\'un spectacle que vous envisagez. Ce feed '
    'intelligent évite le bruit pour ne garder que le signal culturel pertinent.'
)

doc.add_paragraph(
    'Les groupes d\'intérêt permettent de fédérer des communautés autour de passions spécifiques. Le '
    'groupe "Jazz in Paris" rassemble les amateurs de jazz parisiens qui partagent les bonnes adresses, '
    'organisent des sorties communes, et débattent des derniers concerts. "Parents Culturels" aide les '
    'parents à trouver des activités adaptées aux enfants. "Étudiants Fauchés Mais Cultivés" partage '
    'exclusivement les bons plans gratuits ou à prix réduit. Ces groupes deviennent des véritables '
    'communautés avec leurs codes, leurs habitués, leurs événements récurrents.'
)

h2 = doc.add_heading('Organiser des sorties de groupe', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'L\'organisation de sorties culturelles de groupe est souvent un casse-tête logistique que Culture '
    'Radar transforme en plaisir collaboratif. La fonctionnalité "Créer une sortie" lance un processus '
    'guidé où vous définissez d\'abord le type d\'expérience recherchée : sortie entre amis décontractée, '
    'événement d\'entreprise formel, activité familiale, ou date romantique. Selon le contexte, l\'interface '
    's\'adapte pour proposer les options pertinentes.'
)

doc.add_paragraph(
    'L\'invitation des participants peut se faire de multiples façons. Invitez directement vos connexions '
    'Culture Radar, envoyez un lien d\'invitation par email ou SMS aux non-membres, ou créez un événement '
    '"ouvert" que vos amis d\'amis peuvent rejoindre. Le système de RSVP est sophistiqué : au-delà du '
    'simple oui/non, les invités peuvent indiquer leurs contraintes horaires, leurs préférences de budget, '
    'leurs restrictions (accessibilité, régime alimentaire pour les événements avec restauration). '
    'L\'algorithme trouve alors le créneau et l\'événement qui conviennent au maximum de participants.'
)

doc.add_paragraph(
    'Le chat de groupe intégré facilite la coordination. Plus qu\'une simple messagerie, il intègre des '
    'fonctionnalités spécifiques : sondages pour voter sur les options, partage de la localisation en '
    'temps réel le jour J, split de l\'addition si l\'événement est payant, album photo partagé après '
    'l\'événement. L\'historique de ces sorties de groupe constitue progressivement une mémoire collective '
    'précieuse de vos expériences culturelles partagées.'
)

doc.add_paragraph(
    'Pour les organisateurs réguliers - qu\'ils soient professionnels du team building ou simplement les '
    'planificateurs naturels de leur groupe d\'amis - nous offrons des outils avancés. Templates de '
    'sorties réutilisables, gestion de listes d\'invités récurrentes, négociation de tarifs de groupe '
    'avec les lieux partenaires, et même facturation intégrée pour les prestations professionnelles. '
    'Ces power users deviennent souvent des ambassadeurs Culture Radar, créant des expériences mémorables '
    'pour leurs communautés.'
)

# CHAPITRE 10 - FONCTIONNALITÉS PREMIUM
doc.add_page_break()
h1 = doc.add_heading('Chapitre 10 : L\'expérience Premium', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Pourquoi passer à Premium', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Culture Radar Premium n\'est pas simplement une version sans publicité de notre service gratuit. '
    'C\'est une expérience culturelle augmentée, conçue pour les véritables passionnés qui veulent '
    'maximiser leur vie culturelle. À 4,99€ par mois - moins que le prix d\'un café dans la plupart '
    'des lieux culturels - Premium transforme votre façon de découvrir, planifier et vivre la culture. '
    'Laissez-moi vous expliquer pourquoi des milliers d\'utilisateurs considèrent cet investissement '
    'comme indispensable à leur épanouissement culturel.'
)

doc.add_paragraph(
    'L\'algorithme Premium est notre version sans bride de l\'IA de recommandation. Alors que la version '
    'gratuite analyse vos 30 derniers jours d\'activité, Premium examine votre historique complet pour '
    'des recommandations d\'une précision chirurgicale. Il détecte les patterns subtils : votre attirance '
    'pour les événements underground le premier jeudi du mois, votre préférence pour les expositions '
    'photographiques après les périodes stressantes, votre tendance à explorer de nouveaux genres musicaux '
    'en automne. Cette compréhension profonde génère des suggestions qui semblent lire dans vos pensées.'
)

doc.add_paragraph(
    'L\'accès anticipé aux événements exclusifs est un avantage majeur dans les grandes villes où les '
    'meilleures expériences culturelles se vendent en quelques heures. Les membres Premium reçoivent les '
    'annonces 48 heures avant le grand public, avec souvent des préventes exclusives à tarif réduit. '
    'Pour un amateur de théâtre parisien, cela peut faire la différence entre assister à la première '
    'd\'une pièce très attendue ou la manquer complètement. Nous négocions également des places réservées '
    'Premium pour les événements sold-out, accessibles jusqu\'à la dernière minute.'
)

doc.add_paragraph(
    'Les statistiques culturelles personnelles transforment votre activité culturelle en insights '
    'fascinants. Découvrez votre "empreinte culturelle" : combien d\'heures de culture consommées, '
    'diversité des expériences, évolution de vos goûts, comparaison avec votre communauté. Votre '
    '"Spotify Wrapped" culturel en fin d\'année devient un moment de fierté partagée. Ces données '
    'vous aident aussi à équilibrer votre régime culturel : trop de concerts et pas assez d\'expos ? '
    'L\'algorithme ajustera ses recommandations pour vous aider à explorer de nouveaux territoires.'
)

h2 = doc.add_heading('Fonctionnalités exclusives détaillées', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Le mode "Conciergerie Culturelle" est votre assistant personnel haut de gamme. Décrivez vaguement '
    'ce que vous cherchez - "Je veux impressionner des clients japonais en visite la semaine prochaine" - '
    'et recevez un itinéraire culturel sur mesure. Le système prend en compte les spécificités culturelles '
    '(les Japonais apprécient souvent l\'artisanat français), les contraintes pratiques (proximité de '
    'leur hôtel), et même la météo prévue. Cette fonctionnalité seule justifie l\'abonnement pour les '
    'professionnels du tourisme ou de l\'événementiel.'
)

doc.add_paragraph(
    'Les "Parcours Culturels" sont des expériences gamifiées d\'exploration culturelle. Lancez le parcours '
    '"De Montmartre à Belleville : l\'art urbain parisien" et suivez un itinéraire interactif ponctué '
    'd\'événements, d\'anecdotes historiques, et de défis photo. Complétez des parcours pour débloquer '
    'des badges, des réductions exclusives, et même des invitations à des événements privés. C\'est le '
    'parfait mélange entre Pokemon Go et guide culturel expert, transformant votre ville en terrain de '
    'jeu culturel.'
)

doc.add_paragraph(
    'L\'API Premium permet aux power users d\'intégrer Culture Radar dans leur écosystème numérique. '
    'Créez des automatisations complexes : ajoutez automatiquement les concerts de vos artistes Spotify '
    'favoris à vos favoris Culture Radar, synchronisez vos événements avec votre système de notes, '
    'générez des rapports mensuels de votre activité culturelle pour votre blog. Les développeurs de '
    'notre communauté ont créé des intégrations ingénieuses, from affichages muraux connectés montrant '
    'les événements du jour to bots Slack qui suggèrent des team buildings culturels.'
)

doc.add_paragraph(
    'Le support Premium n\'est pas qu\'une hotline prioritaire. C\'est un service de conseil culturel '
    'personnalisé. Nos Cultural Advisors, véritables experts culturels locaux, répondent à vos questions '
    'complexes, vous aident à planifier des occasions spéciales, et peuvent même faire des réservations '
    'en votre nom. "J\'organise un enterrement de vie de garçon culturel pour 8 personnes avec des goûts '
    'très différents" devient un challenge excitant pour nos advisors qui créeront une expérience mémorable.'
)

# CONCLUSION ET RESSOURCES
doc.add_page_break()
h1 = doc.add_heading('Conclusion : Votre voyage culturel commence maintenant', 1)
set_heading_style(h1, 1)

doc.add_paragraph(
    'Vous voici arrivé au terme de ce guide complet de Culture Radar, mais c\'est en réalité le début '
    'de votre aventure culturelle enrichie. À travers ces pages, nous avons exploré ensemble les '
    'multiples facettes de notre plateforme, des fonctionnalités de base aux outils les plus sophistiqués. '
    'Vous avez découvert comment notre intelligence artificielle travaille en coulisses pour personnaliser '
    'votre expérience, comment notre système social transforme la culture en expérience partagée, et '
    'comment nos outils de planification intègrent harmonieusement la culture dans votre quotidien.'
)

doc.add_paragraph(
    'Mais Culture Radar est plus qu\'une somme de fonctionnalités. C\'est une philosophie, une certaine '
    'vision de la culture comme élément essentiel du bien-être et de l\'épanouissement personnel. Dans '
    'un monde où nous sommes constamment sollicités par mille distractions numériques, où le temps libre '
    'est précieux et rare, Culture Radar vous reconnecte avec les expériences réelles, humaines, '
    'transformatrices que seule la culture peut offrir. Chaque concert, chaque exposition, chaque pièce '
    'de théâtre est une opportunité de croissance, de connexion, d\'émerveillement.'
)

doc.add_paragraph(
    'N\'oubliez pas que Culture Radar est un outil évolutif qui s\'améliore constamment grâce à vous. '
    'Vos retours, vos suggestions, vos façons créatives d\'utiliser la plateforme nous inspirent pour '
    'développer de nouvelles fonctionnalités. Rejoignez notre communauté sur forum.culture-radar.fr pour '
    'échanger avec d\'autres passionnés, participer aux beta tests des nouvelles fonctionnalités, et '
    'influencer la direction du produit. Votre voix compte dans la construction de l\'avenir de la '
    'découverte culturelle.'
)

doc.add_paragraph(
    'Enfin, souvenez-vous que la technologie n\'est qu\'un moyen, pas une fin. Culture Radar existe pour '
    'vous libérer du stress de la recherche et de la planification, pour que vous puissiez vous concentrer '
    'sur ce qui compte vraiment : vivre pleinement ces moments culturels. Alors fermez ce manuel, ouvrez '
    'l\'application, et laissez-vous guider vers votre prochaine expérience culturelle mémorable. La '
    'culture vous attend, et grâce à Culture Radar, elle n\'a jamais été aussi accessible.'
)

h2 = doc.add_heading('Ressources et support', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Votre apprentissage ne s\'arrête pas ici. Culture Radar met à votre disposition un écosystème complet '
    'de ressources pour approfondir votre maîtrise de la plateforme. Notre centre d\'aide en ligne '
    '(aide.culture-radar.fr) contient des centaines d\'articles détaillés, constamment mis à jour avec '
    'les dernières fonctionnalités. Les tutoriels vidéo sur notre chaîne YouTube (@CultureRadarFR) offrent '
    'des démonstrations visuelles pour les apprenants visuels. Nos webinaires mensuels gratuits permettent '
    'd\'apprendre directement de notre équipe et de poser vos questions en direct.'
)

doc.add_paragraph(
    'Pour un support personnalisé, plusieurs options s\'offrent à vous. Le chat intégré dans l\'application '
    'connecte avec notre équipe support du lundi au vendredi de 9h à 19h, et le samedi de 10h à 17h. '
    'L\'email support@culture-radar.fr garantit une réponse sous 24-48h pour les questions complexes. '
    'Les membres Premium bénéficient d\'une ligne prioritaire avec des Cultural Advisors dédiés. Pour les '
    'urgences (événement annulé, problème de billetterie), une hotline est disponible 7j/7.'
)

doc.add_paragraph(
    'Restez informé des nouveautés en suivant nos canaux de communication officiels. Notre blog '
    '(blog.culture-radar.fr) publie des articles de fond sur les tendances culturelles, des interviews '
    'd\'artistes, et des guides de ville. La newsletter mensuelle résume les nouvelles fonctionnalités '
    'et partage les meilleures découvertes de la communauté. Nos réseaux sociaux (@CultureRadar sur '
    'toutes les plateformes) offrent de l\'inspiration quotidienne et des concours réguliers.'
)

doc.add_paragraph(
    'Merci de faire partie de la révolution Culture Radar. Ensemble, nous rendons la culture plus '
    'accessible, plus sociale, et plus enrichissante pour tous. Votre voyage culturel ne fait que '
    'commencer, et nous sommes honorés de vous accompagner à chaque étape. Bonne découverte !'
)

# Sauvegarder le document
doc.save('/root/culture-radar/Manuel_Utilisateur_Narratif_Culture_Radar.docx')
print("✅ Manuel utilisateur narratif créé : Manuel_Utilisateur_Narratif_Culture_Radar.docx (30+ pages)")