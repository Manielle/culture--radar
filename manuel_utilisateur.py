#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Créer un nouveau document
doc = Document()

# Fonction pour le style des titres
def set_heading_style(heading, level=1):
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        if level == 1:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(102, 126, 234)  # Bleu Culture Radar
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(118, 75, 162)  # Violet
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True

# PAGE DE GARDE
title = doc.add_heading('CULTURE RADAR', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Manuel d\'Utilisateur')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(24)
subtitle.runs[0].font.bold = True

doc.add_paragraph('')
doc.add_paragraph('')

version = doc.add_paragraph('Version 1.0 - Août 2024')
version.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
doc.add_paragraph('')

tagline = doc.add_paragraph('Votre boussole culturelle personnalisée')
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
tagline.runs[0].font.italic = True
tagline.runs[0].font.size = Pt(14)

# SOMMAIRE
doc.add_page_break()
h1 = doc.add_heading('SOMMAIRE', 1)
set_heading_style(h1, 1)

doc.add_paragraph('1. Bienvenue sur Culture Radar ............................ 3')
doc.add_paragraph('2. Premiers pas .......................................... 4')
doc.add_paragraph('3. Création de compte .................................... 6')
doc.add_paragraph('4. Navigation et interface ............................... 8')
doc.add_paragraph('5. Recherche d\'événements ............................... 10')
doc.add_paragraph('6. Système de recommandations ........................... 12')
doc.add_paragraph('7. Gestion des favoris .................................. 14')
doc.add_paragraph('8. Calendrier personnel ................................. 16')
doc.add_paragraph('9. Notifications ........................................ 18')
doc.add_paragraph('10. Partage social ...................................... 20')
doc.add_paragraph('11. Paramètres du compte ................................ 22')
doc.add_paragraph('12. Abonnement Premium .................................. 24')
doc.add_paragraph('13. Application mobile .................................. 26')
doc.add_paragraph('14. Résolution des problèmes ............................ 28')
doc.add_paragraph('15. FAQ ................................................. 29')
doc.add_paragraph('16. Contact et support .................................. 30')

# CHAPITRE 1 - BIENVENUE
doc.add_page_break()
h1 = doc.add_heading('1. Bienvenue sur Culture Radar', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Qu\'est-ce que Culture Radar ?', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Culture Radar est votre compagnon intelligent pour découvrir les événements culturels '
    'qui vous correspondent vraiment. Notre plateforme utilise l\'intelligence artificielle '
    'pour analyser vos goûts, votre localisation et vos disponibilités afin de vous proposer '
    'des sorties culturelles personnalisées.'
)

doc.add_paragraph('')

h2 = doc.add_heading('Nos valeurs', 2)
set_heading_style(h2, 2)

doc.add_paragraph('• Accessibilité : La culture pour tous, partout')
doc.add_paragraph('• Personnalisation : Des recommandations qui vous ressemblent')
doc.add_paragraph('• Proximité : Valoriser la culture locale')
doc.add_paragraph('• Simplicité : Une interface intuitive et épurée')

h2 = doc.add_heading('Ce que vous pouvez faire avec Culture Radar', 2)
set_heading_style(h2, 2)

doc.add_paragraph('✓ Découvrir des événements près de chez vous')
doc.add_paragraph('✓ Recevoir des recommandations personnalisées')
doc.add_paragraph('✓ Planifier vos sorties culturelles')
doc.add_paragraph('✓ Partager vos découvertes avec vos amis')
doc.add_paragraph('✓ Ne plus jamais manquer un événement qui vous intéresse')

# CHAPITRE 2 - PREMIERS PAS
doc.add_page_break()
h1 = doc.add_heading('2. Premiers pas', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Accéder à Culture Radar', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Culture Radar est accessible depuis :')
doc.add_paragraph('')
doc.add_paragraph('1. Votre navigateur web')
doc.add_paragraph('   Rendez-vous sur www.culture-radar.fr')
doc.add_paragraph('   Compatible avec Chrome, Firefox, Safari, Edge')
doc.add_paragraph('')
doc.add_paragraph('2. L\'application mobile')
doc.add_paragraph('   Téléchargez l\'app sur l\'App Store ou Google Play')
doc.add_paragraph('   Recherchez "Culture Radar"')

h2 = doc.add_heading('Configuration minimale requise', 2)
set_heading_style(h2, 2)

h3 = doc.add_heading('Pour le web', 3)
set_heading_style(h3, 3)

doc.add_paragraph('• Navigateur à jour (versions récentes)')
doc.add_paragraph('• Connexion internet stable')
doc.add_paragraph('• JavaScript activé')
doc.add_paragraph('• Cookies acceptés pour le fonctionnement')

h3 = doc.add_heading('Pour mobile', 3)
set_heading_style(h3, 3)

doc.add_paragraph('• iOS 14+ ou Android 8+')
doc.add_paragraph('• 100 MB d\'espace disponible')
doc.add_paragraph('• Connexion 4G ou Wi-Fi recommandée')

h2 = doc.add_heading('Premier lancement', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Lors de votre première visite, Culture Radar vous accueille avec un tutoriel interactif '
    'de 3 étapes pour personnaliser votre expérience :'
)
doc.add_paragraph('')
doc.add_paragraph('Étape 1 : Choisissez vos centres d\'intérêt')
doc.add_paragraph('Étape 2 : Indiquez votre localisation')
doc.add_paragraph('Étape 3 : Définissez vos préférences de notification')

# CHAPITRE 3 - CRÉATION DE COMPTE
doc.add_page_break()
h1 = doc.add_heading('3. Création de compte', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Pourquoi créer un compte ?', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Un compte Culture Radar vous permet de :')
doc.add_paragraph('• Sauvegarder vos préférences')
doc.add_paragraph('• Accéder à vos favoris sur tous vos appareils')
doc.add_paragraph('• Recevoir des notifications personnalisées')
doc.add_paragraph('• Créer votre calendrier culturel')
doc.add_paragraph('• Partager et commenter des événements')

h2 = doc.add_heading('Processus d\'inscription', 2)
set_heading_style(h2, 2)

doc.add_paragraph('1. Cliquez sur "S\'inscrire" en haut à droite')
doc.add_paragraph('2. Choisissez votre méthode d\'inscription :')
doc.add_paragraph('   • Email et mot de passe')
doc.add_paragraph('   • Connexion avec Google')
doc.add_paragraph('   • Connexion avec Facebook')
doc.add_paragraph('   • Connexion avec Apple ID')
doc.add_paragraph('')
doc.add_paragraph('3. Complétez votre profil :')
doc.add_paragraph('   • Prénom (obligatoire)')
doc.add_paragraph('   • Date de naissance (pour les recommandations)')
doc.add_paragraph('   • Code postal (pour la géolocalisation)')
doc.add_paragraph('')
doc.add_paragraph('4. Validez votre email')
doc.add_paragraph('   Un email de confirmation vous sera envoyé')

h2 = doc.add_heading('Sécurité du mot de passe', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Votre mot de passe doit contenir :')
doc.add_paragraph('✓ Au moins 8 caractères')
doc.add_paragraph('✓ Une lettre majuscule')
doc.add_paragraph('✓ Une lettre minuscule')
doc.add_paragraph('✓ Un chiffre')
doc.add_paragraph('✓ Un caractère spécial (!, @, #, $, etc.)')

h3 = doc.add_heading('Conseil de sécurité', 3)
set_heading_style(h3, 3)

doc.add_paragraph(
    'Utilisez un mot de passe unique pour Culture Radar. '
    'Activez l\'authentification à deux facteurs dans les paramètres pour une sécurité maximale.'
)

# CHAPITRE 4 - NAVIGATION
doc.add_page_break()
h1 = doc.add_heading('4. Navigation et interface', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Structure de l\'interface', 2)
set_heading_style(h2, 2)

h3 = doc.add_heading('Barre de navigation principale', 3)
set_heading_style(h3, 3)

doc.add_paragraph('La barre de navigation en haut de l\'écran contient :')
doc.add_paragraph('')
doc.add_paragraph('• Logo Culture Radar (retour à l\'accueil)')
doc.add_paragraph('• Barre de recherche centrale')
doc.add_paragraph('• Icône Favoris (cœur)')
doc.add_paragraph('• Icône Notifications (cloche)')
doc.add_paragraph('• Menu utilisateur (avatar)')

h3 = doc.add_heading('Menu latéral', 3)
set_heading_style(h3, 3)

doc.add_paragraph('Le menu latéral gauche propose :')
doc.add_paragraph('')
doc.add_paragraph('📍 Découvrir : Explorer les événements')
doc.add_paragraph('📅 Calendrier : Vos événements planifiés')
doc.add_paragraph('⭐ Favoris : Vos événements sauvegardés')
doc.add_paragraph('🎯 Pour vous : Recommandations personnalisées')
doc.add_paragraph('🗺️ Carte : Vue géographique des événements')
doc.add_paragraph('⚙️ Paramètres : Configuration du compte')

h2 = doc.add_heading('Zone de contenu principal', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'La zone centrale affiche le contenu principal selon votre navigation. '
    'Les événements sont présentés sous forme de cartes interactives avec :'
)
doc.add_paragraph('')
doc.add_paragraph('• Photo de l\'événement')
doc.add_paragraph('• Titre et description courte')
doc.add_paragraph('• Date et horaire')
doc.add_paragraph('• Lieu et distance')
doc.add_paragraph('• Prix ou mention "Gratuit"')
doc.add_paragraph('• Boutons d\'action rapide')

h2 = doc.add_heading('Filtres latéraux', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Sur la droite, accédez aux filtres :')
doc.add_paragraph('• Catégories (concerts, théâtre, expos...)')
doc.add_paragraph('• Date (aujourd\'hui, week-end, semaine...)')
doc.add_paragraph('• Distance (1km, 5km, 10km...)')
doc.add_paragraph('• Prix (gratuit, moins de 20€...)')
doc.add_paragraph('• Accessibilité PMR')

# CHAPITRE 5 - RECHERCHE
doc.add_page_break()
h1 = doc.add_heading('5. Recherche d\'événements', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Barre de recherche intelligente', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'La barre de recherche principale comprend et interprète vos requêtes naturelles :'
)
doc.add_paragraph('')
doc.add_paragraph('Exemples de recherches :')
doc.add_paragraph('• "Concert jazz ce soir"')
doc.add_paragraph('• "Expo gratuite Paris 11"')
doc.add_paragraph('• "Théâtre dimanche après-midi"')
doc.add_paragraph('• "Activités enfants vacances"')

h2 = doc.add_heading('Recherche avancée', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Cliquez sur "Recherche avancée" pour accéder à :')
doc.add_paragraph('')

h3 = doc.add_heading('Critères de base', 3)
set_heading_style(h3, 3)

doc.add_paragraph('• Mots-clés spécifiques')
doc.add_paragraph('• Nom de l\'artiste ou du lieu')
doc.add_paragraph('• Type d\'événement')
doc.add_paragraph('• Période précise')

h3 = doc.add_heading('Critères avancés', 3)
set_heading_style(h3, 3)

doc.add_paragraph('• Durée de l\'événement')
doc.add_paragraph('• Langue de l\'événement')
doc.add_paragraph('• Âge recommandé')
doc.add_paragraph('• Jauge (petit/moyen/grand événement)')
doc.add_paragraph('• Labels (éco-responsable, accessible PMR...)')

h2 = doc.add_heading('Historique de recherche', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Culture Radar mémorise vos 10 dernières recherches pour un accès rapide. '
    'Vous pouvez effacer cet historique dans les paramètres de confidentialité.'
)

h2 = doc.add_heading('Recherche vocale', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Sur mobile et navigateurs compatibles, utilisez le micro pour dicter votre recherche. '
    'Dites par exemple : "Trouve-moi une pièce de théâtre pour demain soir".'
)

# CHAPITRE 6 - RECOMMANDATIONS
doc.add_page_break()
h1 = doc.add_heading('6. Système de recommandations', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Comment fonctionne notre IA ?', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Notre algorithme analyse :')
doc.add_paragraph('')
doc.add_paragraph('1. Vos préférences déclarées')
doc.add_paragraph('   • Catégories favorites')
doc.add_paragraph('   • Artistes suivis')
doc.add_paragraph('   • Lieux préférés')
doc.add_paragraph('')
doc.add_paragraph('2. Votre comportement')
doc.add_paragraph('   • Événements consultés')
doc.add_paragraph('   • Temps passé sur les fiches')
doc.add_paragraph('   • Événements ajoutés aux favoris')
doc.add_paragraph('')
doc.add_paragraph('3. Votre contexte')
doc.add_paragraph('   • Localisation actuelle')
doc.add_paragraph('   • Météo du jour')
doc.add_paragraph('   • Jour de la semaine')
doc.add_paragraph('   • Événements déjà planifiés')

h2 = doc.add_heading('Types de recommandations', 2)
set_heading_style(h2, 2)

h3 = doc.add_heading('Découvertes du jour', 3)
set_heading_style(h3, 3)

doc.add_paragraph(
    'Chaque matin, 5 événements sélectionnés spécialement pour vous, '
    'basés sur vos goûts et la programmation du jour.'
)

h3 = doc.add_heading('Similaires', 3)
set_heading_style(h3, 3)

doc.add_paragraph(
    'Quand vous consultez un événement, découvrez des suggestions similaires '
    'en bas de page : même style, même lieu, même période.'
)

h3 = doc.add_heading('Tendances locales', 3)
set_heading_style(h3, 3)

doc.add_paragraph(
    'Les événements populaires dans votre quartier, plébiscités par des utilisateurs '
    'ayant des goûts similaires aux vôtres.'
)

h2 = doc.add_heading('Améliorer les recommandations', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Pour des suggestions plus pertinentes :')
doc.add_paragraph('✓ Notez les événements auxquels vous avez assisté')
doc.add_paragraph('✓ Utilisez les boutons "J\'aime" et "Pas intéressé"')
doc.add_paragraph('✓ Complétez votre profil culturel')
doc.add_paragraph('✓ Suivez vos artistes et lieux favoris')

# CHAPITRE 7 - FAVORIS
doc.add_page_break()
h1 = doc.add_heading('7. Gestion des favoris', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Ajouter aux favoris', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Pour sauvegarder un événement :')
doc.add_paragraph('1. Cliquez sur l\'icône cœur sur la carte événement')
doc.add_paragraph('2. Ou ouvrez la fiche et cliquez sur "Ajouter aux favoris"')
doc.add_paragraph('')
doc.add_paragraph('L\'événement est immédiatement sauvegardé dans votre liste.')

h2 = doc.add_heading('Organiser vos favoris', 2)
set_heading_style(h2, 2)

h3 = doc.add_heading('Créer des listes', 3)
set_heading_style(h3, 3)

doc.add_paragraph('Organisez vos favoris en listes thématiques :')
doc.add_paragraph('• Sorties en famille')
doc.add_paragraph('• Dates romantiques')
doc.add_paragraph('• Avec les amis')
doc.add_paragraph('• À voir absolument')
doc.add_paragraph('• Idées cadeaux')

h3 = doc.add_heading('Gérer les listes', 3)
set_heading_style(h3, 3)

doc.add_paragraph('Pour chaque liste, vous pouvez :')
doc.add_paragraph('• Renommer la liste')
doc.add_paragraph('• Ajouter une description')
doc.add_paragraph('• Choisir une couleur')
doc.add_paragraph('• Définir la confidentialité (privée/publique)')
doc.add_paragraph('• Partager avec des amis')

h2 = doc.add_heading('Notifications de favoris', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Recevez des alertes pour vos événements favoris :')
doc.add_paragraph('• Ouverture de la billetterie')
doc.add_paragraph('• Dernières places disponibles')
doc.add_paragraph('• Changement de date ou d\'horaire')
doc.add_paragraph('• Promotion sur les billets')

h2 = doc.add_heading('Export des favoris', 2)
set_heading_style(h2, 2)

doc.add_paragraph(
    'Exportez votre liste de favoris en PDF ou ajoutez-les '
    'directement à votre calendrier (Google, Apple, Outlook).'
)

# CHAPITRE 8 - CALENDRIER
doc.add_page_break()
h1 = doc.add_heading('8. Calendrier personnel', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Vue d\'ensemble', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Le calendrier Culture Radar propose 3 vues :')
doc.add_paragraph('')
doc.add_paragraph('• Vue Mois : Vision globale de vos sorties')
doc.add_paragraph('• Vue Semaine : Planning détaillé')
doc.add_paragraph('• Vue Liste : Tous vos événements chronologiques')

h2 = doc.add_heading('Ajouter un événement', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Depuis une fiche événement :')
doc.add_paragraph('1. Cliquez sur "Ajouter au calendrier"')
doc.add_paragraph('2. Choisissez la date (si plusieurs dates)')
doc.add_paragraph('3. Ajoutez un rappel (optionnel)')
doc.add_paragraph('4. Invitez des amis (optionnel)')

h2 = doc.add_heading('Rappels et notifications', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Configurez vos rappels :')
doc.add_paragraph('• 1 semaine avant')
doc.add_paragraph('• 3 jours avant')
doc.add_paragraph('• 1 jour avant')
doc.add_paragraph('• Le matin même')
doc.add_paragraph('• 2 heures avant')

h2 = doc.add_heading('Synchronisation', 2)
set_heading_style(h2, 2)

h3 = doc.add_heading('Import', 3)
set_heading_style(h3, 3)

doc.add_paragraph('Importez vos événements depuis :')
doc.add_paragraph('• Google Calendar')
doc.add_paragraph('• Apple Calendar')
doc.add_paragraph('• Outlook')
doc.add_paragraph('• Fichier .ics')

h3 = doc.add_heading('Export', 3)
set_heading_style(h3, 3)

doc.add_paragraph('Exportez votre calendrier Culture Radar vers vos apps préférées '
                 'pour une vision unifiée de votre agenda.')

h2 = doc.add_heading('Partage du calendrier', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Partagez votre calendrier culturel :')
doc.add_paragraph('• Lien public (lecture seule)')
doc.add_paragraph('• Invitation privée (amis peuvent suggérer)')
doc.add_paragraph('• Intégration sur votre blog/site')

# CHAPITRE 9 - NOTIFICATIONS
doc.add_page_break()
h1 = doc.add_heading('9. Notifications', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Types de notifications', 2)
set_heading_style(h2, 2)

h3 = doc.add_heading('Recommandations', 3)
set_heading_style(h3, 3)

doc.add_paragraph('• Nouveaux événements correspondant à vos goûts')
doc.add_paragraph('• Événements populaires près de vous')
doc.add_paragraph('• Dernière chance (événements se terminant bientôt)')

h3 = doc.add_heading('Rappels', 3)
set_heading_style(h3, 3)

doc.add_paragraph('• Événements de votre calendrier')
doc.add_paragraph('• Ouverture de billetterie')
doc.add_paragraph('• Événements favoris approchant')

h3 = doc.add_heading('Social', 3)
set_heading_style(h3, 3)

doc.add_paragraph('• Invitations d\'amis')
doc.add_paragraph('• Événements partagés avec vous')
doc.add_paragraph('• Nouveaux followers')

h3 = doc.add_heading('Actualités', 3)
set_heading_style(h3, 3)

doc.add_paragraph('• Annonces de vos lieux suivis')
doc.add_paragraph('• Nouvelles dates de vos artistes favoris')
doc.add_paragraph('• Offres spéciales et promotions')

h2 = doc.add_heading('Paramétrer les notifications', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Dans Paramètres > Notifications :')
doc.add_paragraph('')
doc.add_paragraph('1. Choisissez les types à recevoir')
doc.add_paragraph('2. Définissez la fréquence :')
doc.add_paragraph('   • Temps réel')
doc.add_paragraph('   • Résumé quotidien')
doc.add_paragraph('   • Résumé hebdomadaire')
doc.add_paragraph('3. Sélectionnez les canaux :')
doc.add_paragraph('   • Notifications push (mobile)')
doc.add_paragraph('   • Email')
doc.add_paragraph('   • SMS (Premium)')

h2 = doc.add_heading('Mode Ne pas déranger', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Activez le mode silencieux :')
doc.add_paragraph('• Pendant vos heures de travail')
doc.add_paragraph('• La nuit (22h-8h par défaut)')
doc.add_paragraph('• En vacances')

# CHAPITRE 10 - PARTAGE SOCIAL
doc.add_page_break()
h1 = doc.add_heading('10. Partage social', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Partager un événement', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Depuis la fiche événement, partagez via :')
doc.add_paragraph('')
doc.add_paragraph('• Message direct Culture Radar')
doc.add_paragraph('• WhatsApp')
doc.add_paragraph('• Facebook')
doc.add_paragraph('• Twitter')
doc.add_paragraph('• Instagram Stories')
doc.add_paragraph('• Email')
doc.add_paragraph('• Copier le lien')

h2 = doc.add_heading('Inviter des amis', 2)
set_heading_style(h2, 2)

h3 = doc.add_heading('Créer un groupe', 3)
set_heading_style(h3, 3)

doc.add_paragraph('1. Sur l\'événement, cliquez "Inviter des amis"')
doc.add_paragraph('2. Sélectionnez vos contacts Culture Radar')
doc.add_paragraph('3. Ou invitez par email/SMS')
doc.add_paragraph('4. Créez un groupe de discussion')

h3 = doc.add_heading('Gérer les réponses', 3)
set_heading_style(h3, 3)

doc.add_paragraph('Visualisez qui :')
doc.add_paragraph('• Participe ✓')
doc.add_paragraph('• Peut-être ?')
doc.add_paragraph('• Ne participe pas ✗')
doc.add_paragraph('• N\'a pas encore répondu ○')

h2 = doc.add_heading('Communauté Culture Radar', 2)
set_heading_style(h2, 2)

h3 = doc.add_heading('Suivre des utilisateurs', 3)
set_heading_style(h3, 3)

doc.add_paragraph('Suivez d\'autres passionnés pour découvrir :')
doc.add_paragraph('• Leurs événements favoris')
doc.add_paragraph('• Leurs listes publiques')
doc.add_paragraph('• Leurs avis et notes')

h3 = doc.add_heading('Avis et commentaires', 3)
set_heading_style(h3, 3)

doc.add_paragraph('Après un événement :')
doc.add_paragraph('• Notez de 1 à 5 étoiles')
doc.add_paragraph('• Laissez un commentaire')
doc.add_paragraph('• Ajoutez des photos')
doc.add_paragraph('• Recommandez à vos followers')

# CHAPITRE 11 - PARAMÈTRES
doc.add_page_break()
h1 = doc.add_heading('11. Paramètres du compte', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Profil personnel', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Modifiez vos informations :')
doc.add_paragraph('• Photo de profil')
doc.add_paragraph('• Nom d\'affichage')
doc.add_paragraph('• Bio (140 caractères)')
doc.add_paragraph('• Ville de résidence')
doc.add_paragraph('• Langues parlées')

h2 = doc.add_heading('Préférences culturelles', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Affinez vos goûts :')
doc.add_paragraph('')

h3 = doc.add_heading('Catégories favorites', 3)
set_heading_style(h3, 3)

doc.add_paragraph('Notez de 1 à 5 votre intérêt pour :')
doc.add_paragraph('• Concerts & Musique')
doc.add_paragraph('• Théâtre & Spectacles')
doc.add_paragraph('• Expositions & Musées')
doc.add_paragraph('• Cinéma')
doc.add_paragraph('• Conférences & Débats')
doc.add_paragraph('• Ateliers & Cours')
doc.add_paragraph('• Festivals')
doc.add_paragraph('• Sport & Loisirs')

h3 = doc.add_heading('Critères de sélection', 3)
set_heading_style(h3, 3)

doc.add_paragraph('• Budget maximum par sortie')
doc.add_paragraph('• Distance maximale acceptable')
doc.add_paragraph('• Jours de sortie préférés')
doc.add_paragraph('• Horaires préférés')

h2 = doc.add_heading('Confidentialité', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Contrôlez vos données :')
doc.add_paragraph('• Profil public/privé')
doc.add_paragraph('• Géolocalisation (toujours/jamais/demander)')
doc.add_paragraph('• Partage avec partenaires (opt-in/opt-out)')
doc.add_paragraph('• Télécharger mes données')
doc.add_paragraph('• Supprimer mon compte')

h2 = doc.add_heading('Sécurité', 2)
set_heading_style(h2, 2)

doc.add_paragraph('• Changer le mot de passe')
doc.add_paragraph('• Authentification à deux facteurs')
doc.add_paragraph('• Sessions actives')
doc.add_paragraph('• Historique de connexion')

# CHAPITRE 12 - PREMIUM
doc.add_page_break()
h1 = doc.add_heading('12. Abonnement Premium', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Avantages Premium', 2)
set_heading_style(h2, 2)

doc.add_paragraph('L\'abonnement Premium offre :')
doc.add_paragraph('')
doc.add_paragraph('✨ Recommandations illimitées')
doc.add_paragraph('✨ Alertes prioritaires sur les nouveautés')
doc.add_paragraph('✨ Réservation directe intégrée')
doc.add_paragraph('✨ Accès aux événements exclusifs')
doc.add_paragraph('✨ Statistiques de vos sorties culturelles')
doc.add_paragraph('✨ Export illimité en PDF')
doc.add_paragraph('✨ Support prioritaire')
doc.add_paragraph('✨ Sans publicité')

h2 = doc.add_heading('Tarifs', 2)
set_heading_style(h2, 2)

doc.add_paragraph('• Mensuel : 4,99€/mois')
doc.add_paragraph('• Annuel : 49,99€/an (2 mois offerts)')
doc.add_paragraph('• Étudiant : 2,99€/mois (justificatif requis)')
doc.add_paragraph('• Famille : 7,99€/mois (jusqu\'à 4 comptes)')

h2 = doc.add_heading('Souscrire à Premium', 2)
set_heading_style(h2, 2)

doc.add_paragraph('1. Allez dans Paramètres > Abonnement')
doc.add_paragraph('2. Choisissez votre formule')
doc.add_paragraph('3. Entrez vos informations de paiement')
doc.add_paragraph('4. Validez (activation immédiate)')

h3 = doc.add_heading('Moyens de paiement acceptés', 3)
set_heading_style(h3, 3)

doc.add_paragraph('• Carte bancaire (Visa, Mastercard, Amex)')
doc.add_paragraph('• PayPal')
doc.add_paragraph('• Apple Pay')
doc.add_paragraph('• Google Pay')

h2 = doc.add_heading('Gérer son abonnement', 2)
set_heading_style(h2, 2)

doc.add_paragraph('• Changer de formule à tout moment')
doc.add_paragraph('• Mettre en pause (jusqu\'à 3 mois)')
doc.add_paragraph('• Résilier sans frais')
doc.add_paragraph('• Télécharger les factures')

h3 = doc.add_heading('Essai gratuit', 3)
set_heading_style(h3, 3)

doc.add_paragraph('Profitez de 30 jours d\'essai gratuit sans engagement. '
                 'Aucune carte bancaire requise pour l\'essai.')

# CHAPITRE 13 - APPLICATION MOBILE
doc.add_page_break()
h1 = doc.add_heading('13. Application mobile', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Téléchargement', 2)
set_heading_style(h2, 2)

h3 = doc.add_heading('iOS (iPhone/iPad)', 3)
set_heading_style(h3, 3)

doc.add_paragraph('1. Ouvrez l\'App Store')
doc.add_paragraph('2. Recherchez "Culture Radar"')
doc.add_paragraph('3. Tapez sur "Obtenir"')
doc.add_paragraph('4. Authentifiez-vous avec Face ID/Touch ID')

h3 = doc.add_heading('Android', 3)
set_heading_style(h3, 3)

doc.add_paragraph('1. Ouvrez Google Play Store')
doc.add_paragraph('2. Recherchez "Culture Radar"')
doc.add_paragraph('3. Tapez sur "Installer"')
doc.add_paragraph('4. Acceptez les permissions')

h2 = doc.add_heading('Fonctionnalités mobiles exclusives', 2)
set_heading_style(h2, 2)

doc.add_paragraph('• Notifications push en temps réel')
doc.add_paragraph('• Géolocalisation précise')
doc.add_paragraph('• Mode hors ligne (favoris et calendrier)')
doc.add_paragraph('• Scanner de QR code pour événements')
doc.add_paragraph('• Partage rapide vers stories')
doc.add_paragraph('• Widget écran d\'accueil')

h2 = doc.add_heading('Permissions requises', 2)
set_heading_style(h2, 2)

doc.add_paragraph('L\'app demande accès à :')
doc.add_paragraph('')
doc.add_paragraph('📍 Localisation : Pour les événements proches')
doc.add_paragraph('📷 Appareil photo : Pour scanner les QR codes')
doc.add_paragraph('📅 Calendrier : Pour synchroniser vos événements')
doc.add_paragraph('🔔 Notifications : Pour les alertes')
doc.add_paragraph('📁 Stockage : Pour sauvegarder les favoris hors ligne')

h3 = doc.add_heading('Note sur la confidentialité', 3)
set_heading_style(h3, 3)

doc.add_paragraph('Toutes les permissions sont optionnelles. L\'app fonctionne '
                 'sans, mais avec des fonctionnalités réduites.')

h2 = doc.add_heading('Synchronisation', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Vos données sont synchronisées automatiquement entre :')
doc.add_paragraph('• Application mobile')
doc.add_paragraph('• Site web')
doc.add_paragraph('• Application tablette')

# CHAPITRE 14 - RÉSOLUTION PROBLÈMES
doc.add_page_break()
h1 = doc.add_heading('14. Résolution des problèmes', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Problèmes de connexion', 2)
set_heading_style(h2, 2)

h3 = doc.add_heading('Mot de passe oublié', 3)
set_heading_style(h3, 3)

doc.add_paragraph('1. Cliquez sur "Mot de passe oublié ?"')
doc.add_paragraph('2. Entrez votre email')
doc.add_paragraph('3. Consultez vos emails (vérifiez les spams)')
doc.add_paragraph('4. Cliquez sur le lien de réinitialisation')
doc.add_paragraph('5. Créez un nouveau mot de passe')

h3 = doc.add_heading('Compte bloqué', 3)
set_heading_style(h3, 3)

doc.add_paragraph('Après 5 tentatives échouées :')
doc.add_paragraph('• Attendez 15 minutes')
doc.add_paragraph('• Ou réinitialisez votre mot de passe')
doc.add_paragraph('• Contactez le support si persistant')

h2 = doc.add_heading('Problèmes d\'affichage', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Si le site ne s\'affiche pas correctement :')
doc.add_paragraph('')
doc.add_paragraph('✓ Videz le cache du navigateur')
doc.add_paragraph('✓ Désactivez les extensions (AdBlock...)')
doc.add_paragraph('✓ Mettez à jour votre navigateur')
doc.add_paragraph('✓ Essayez un autre navigateur')
doc.add_paragraph('✓ Vérifiez votre connexion internet')

h2 = doc.add_heading('Problèmes de géolocalisation', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Si votre position n\'est pas détectée :')
doc.add_paragraph('')
doc.add_paragraph('Sur ordinateur :')
doc.add_paragraph('• Autorisez la géolocalisation dans le navigateur')
doc.add_paragraph('• Vérifiez les paramètres de confidentialité')
doc.add_paragraph('')
doc.add_paragraph('Sur mobile :')
doc.add_paragraph('• Activez le GPS')
doc.add_paragraph('• Autorisez l\'app à accéder à la position')
doc.add_paragraph('• Vérifiez les paramètres iOS/Android')

h2 = doc.add_heading('Notifications non reçues', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Vérifiez :')
doc.add_paragraph('• Les paramètres de notification dans l\'app')
doc.add_paragraph('• Les autorisations système')
doc.add_paragraph('• Que l\'email n\'est pas en spam')
doc.add_paragraph('• Le mode Ne pas déranger')

# CHAPITRE 15 - FAQ
doc.add_page_break()
h1 = doc.add_heading('15. FAQ - Questions fréquentes', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Général', 2)
set_heading_style(h2, 2)

p = doc.add_paragraph()
p.add_run('Q: Culture Radar est-il gratuit ?').bold = True
doc.add_paragraph('R: Oui, l\'utilisation de base est gratuite. Un abonnement Premium '
                 'offre des fonctionnalités avancées.')

p = doc.add_paragraph()
p.add_run('Q: Dans quelles villes fonctionne Culture Radar ?').bold = True
doc.add_paragraph('R: Nous couvrons toute la France, avec un focus sur les grandes '
                 'métropoles : Paris, Lyon, Marseille, Toulouse, Nice, Nantes, '
                 'Strasbourg, Montpellier, Bordeaux, Lille.')

p = doc.add_paragraph()
p.add_run('Q: D\'où viennent les événements ?').bold = True
doc.add_paragraph('R: Nous agrégeons les données de multiples sources : organisateurs '
                 'partenaires, API publiques, OpenAgenda, et contributions communautaires '
                 'vérifiées.')

h2 = doc.add_heading('Compte', 2)
set_heading_style(h2, 2)

p = doc.add_paragraph()
p.add_run('Q: Puis-je utiliser Culture Radar sans compte ?').bold = True
doc.add_paragraph('R: Oui, mais avec des fonctionnalités limitées. Un compte gratuit '
                 'permet d\'accéder aux favoris, calendrier et recommandations.')

p = doc.add_paragraph()
p.add_run('Q: Comment supprimer mon compte ?').bold = True
doc.add_paragraph('R: Dans Paramètres > Confidentialité > Supprimer mon compte. '
                 'La suppression est définitive après 30 jours.')

h2 = doc.add_heading('Technique', 2)
set_heading_style(h2, 2)

p = doc.add_paragraph()
p.add_run('Q: L\'app consomme-t-elle beaucoup de données ?').bold = True
doc.add_paragraph('R: Non, environ 10-20 MB par mois en usage normal. Les images '
                 'sont optimisées et mises en cache.')

p = doc.add_paragraph()
p.add_run('Q: Mes données sont-elles sécurisées ?').bold = True
doc.add_paragraph('R: Oui, nous utilisons un chiffrement de bout en bout et respectons '
                 'le RGPD. Vos données ne sont jamais vendues.')

# CHAPITRE 16 - CONTACT
doc.add_page_break()
h1 = doc.add_heading('16. Contact et support', 1)
set_heading_style(h1, 1)

h2 = doc.add_heading('Centre d\'aide', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Visitez notre centre d\'aide en ligne :')
doc.add_paragraph('🌐 aide.culture-radar.fr')
doc.add_paragraph('')
doc.add_paragraph('Vous y trouverez :')
doc.add_paragraph('• Tutoriels vidéo')
doc.add_paragraph('• Articles détaillés')
doc.add_paragraph('• Mises à jour et nouveautés')
doc.add_paragraph('• État des services')

h2 = doc.add_heading('Contact direct', 2)
set_heading_style(h2, 2)

h3 = doc.add_heading('Email', 3)
set_heading_style(h3, 3)

doc.add_paragraph('📧 support@culture-radar.fr')
doc.add_paragraph('Réponse sous 24-48h en semaine')

h3 = doc.add_heading('Chat en ligne', 3)
set_heading_style(h3, 3)

doc.add_paragraph('💬 Disponible pour les utilisateurs Premium')
doc.add_paragraph('Lun-Ven : 9h-19h')
doc.add_paragraph('Sam : 10h-17h')

h3 = doc.add_heading('Réseaux sociaux', 3)
set_heading_style(h3, 3)

doc.add_paragraph('📱 Twitter : @CultureRadar')
doc.add_paragraph('📘 Facebook : /CultureRadarFR')
doc.add_paragraph('📷 Instagram : @culture_radar')

h2 = doc.add_heading('Signaler un problème', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Pour signaler :')
doc.add_paragraph('• Un bug technique : bug@culture-radar.fr')
doc.add_paragraph('• Un contenu inapproprié : moderation@culture-radar.fr')
doc.add_paragraph('• Un problème de sécurité : security@culture-radar.fr')

h2 = doc.add_heading('Feedback et suggestions', 2)
set_heading_style(h2, 2)

doc.add_paragraph('Vos idées nous intéressent !')
doc.add_paragraph('Envoyez vos suggestions à : idees@culture-radar.fr')
doc.add_paragraph('')
doc.add_paragraph('Ou participez à notre communauté :')
doc.add_paragraph('• Forum : forum.culture-radar.fr')
doc.add_paragraph('• Discord : discord.gg/cultureradar')

# MENTIONS LÉGALES
doc.add_page_break()
doc.add_paragraph('')
doc.add_paragraph('')
doc.add_paragraph('')

legal = doc.add_paragraph('Culture Radar® est une marque déposée de CultureRadar SAS')
legal.alignment = WD_ALIGN_PARAGRAPH.CENTER
legal.runs[0].font.size = Pt(8)

doc.add_paragraph('')

copyright = doc.add_paragraph('© 2024 CultureRadar SAS - Tous droits réservés')
copyright.alignment = WD_ALIGN_PARAGRAPH.CENTER
copyright.runs[0].font.size = Pt(8)

doc.add_paragraph('')

version_info = doc.add_paragraph('Version 1.0 - Août 2024')
version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
version_info.runs[0].font.size = Pt(8)

# Sauvegarder le document
doc.save('/root/culture-radar/Manuel_Utilisateur_Culture_Radar.docx')
print("✅ Manuel utilisateur créé : Manuel_Utilisateur_Culture_Radar.docx (30 pages)")